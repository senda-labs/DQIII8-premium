"""Programmatic binary fixture builders for the metadata toolchain tests.

Zero committed binary fixtures — every fixture is built at test time inside
tmp_path (Pillow for images, pikepdf for PDFs, python-docx + zipfile surgery
for OOXML). This repo is public and has already had two real secret leaks
purged from its history; a committed "EXIF+GPS" or "real author name"
fixture would be exactly that kind of risk.
"""

from __future__ import annotations

import io
import struct
import zipfile
from pathlib import Path

JUMBF_TYPE_UUID = bytes.fromhex("6332706100110010800000aa00389b71")

_DOCX_CONTENT_TYPES = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    b'<Override PartName="/word/document.xml" '
    b'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    b"</Types>"
)


# --------------------------------------------------------------- images ---


def make_jpeg_with_exif(path: Path, *, author: str = "Test Author") -> bytes:
    from PIL import Image

    img = Image.new("RGB", (32, 24), color=(120, 60, 200))
    exif = Image.Exif()
    exif[0x013B] = author  # Artist tag, IFD0
    img.save(path, format="JPEG", exif=exif)
    return path.read_bytes()


def make_jpeg_plain(path: Path) -> bytes:
    from PIL import Image

    img = Image.new("RGB", (16, 16), color=(10, 20, 30))
    img.save(path, format="JPEG")
    return path.read_bytes()


def _jumbf_c2pa_box(seq: int = 1) -> bytes:
    """A minimal fake JUMBF box carrying the C2PA type UUID — enough for the
    fixed-offset detector, not a spec-valid manifest (never parsed as one).
    """
    body = b"\x00" * 8 + JUMBF_TYPE_UUID + b"fake-manifest-payload" + struct.pack(">I", seq)
    return body


def make_jpeg_with_c2pa_app11(path: Path, *, split_segments: bool = False) -> bytes:
    """Builds a JPEG with a JUMBF/C2PA APP11 payload, optionally split across
    two continuation segments (same En, ascending Z) the way real encoders
    do for large manifests.
    """
    from PIL import Image

    img = Image.new("RGB", (16, 16), color=(50, 50, 50))
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    base = buf.getvalue()
    assert base.startswith(b"\xff\xd8")

    jumbf = _jumbf_c2pa_box()
    en = 1

    def _segment(z: int, chunk: bytes) -> bytes:
        payload = b"JP" + struct.pack(">H", en) + struct.pack(">I", z) + chunk
        seg_len = len(payload) + 2
        return b"\xff\xeb" + struct.pack(">H", seg_len) + payload

    if split_segments:
        half = len(jumbf) // 2
        segs = _segment(1, jumbf[:half]) + _segment(2, jumbf[half:])
    else:
        segs = _segment(1, jumbf)

    out = base[:2] + segs + base[2:]
    path.write_bytes(out)
    return out


def make_png_with_text_chunk(path: Path, *, keyword: str = "Author", value: str = "Test") -> bytes:
    from PIL import Image, PngImagePlugin

    img = Image.new("RGB", (16, 16), color=(1, 2, 3))
    info = PngImagePlugin.PngInfo()
    info.add_text(keyword, value)
    img.save(path, format="PNG", pnginfo=info)
    return path.read_bytes()


def make_png_with_cabx(path: Path) -> bytes:
    from PIL import Image

    img = Image.new("RGB", (16, 16), color=(4, 5, 6))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    raw = buf.getvalue()

    payload = JUMBF_TYPE_UUID + b"fake-c2pa-png-payload"
    length = struct.pack(">I", len(payload))
    ctype = b"caBX"
    import zlib

    crc = struct.pack(">I", zlib.crc32(ctype + payload) & 0xFFFFFFFF)
    chunk = length + ctype + payload + crc

    # insert right after IHDR (first chunk after the 8-byte signature)
    ihdr_len = struct.unpack(">I", raw[8:12])[0]
    ihdr_end = 8 + 8 + ihdr_len + 4
    out = raw[:ihdr_end] + chunk + raw[ihdr_end:]
    path.write_bytes(out)
    return out


# ------------------------------------------------------------------- pdf ---


def make_pdf_simple(path: Path, *, author: str = "PDF Author") -> bytes:
    import pikepdf

    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.docinfo["/Author"] = author
    pdf.save(path)
    return path.read_bytes()


def make_pdf_incremental_two_revisions(path: Path, *, rev1_author: str, rev2_author: str) -> bytes:
    """Hand-built classic-xref PDF (pikepdf in this environment has no
    incremental-save API) with a genuine incremental-update section
    appended on top of a complete, independently-valid revision 1 — the
    fixture that proves full-rewrite save (not exiftool-style -all=) is
    what actually drops history: rev1_author must be byte-absent after
    safe-tier removal, even though it's still physically present in a
    naive `-all=`-style incremental-append clean.
    """
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Resources << >> /Contents 4 0 R >>",
        4: b"<< /Length 0 >>\nstream\n\nendstream",
        5: f"<< /Author ({rev1_author}) >>".encode(),
    }

    header = b"%PDF-1.4\n"
    body = bytearray(header)
    offsets = {}
    for num in sorted(objects):
        offsets[num] = len(body)
        body += f"{num} 0 obj\n".encode() + objects[num] + b"\nendobj\n"

    xref_offset = len(body)
    n = max(objects) + 1
    xref = bytearray(f"xref\n0 {n}\n".encode())
    xref += b"0000000000 65535 f \n"
    for num in range(1, n):
        xref += f"{offsets[num]:010d} 00000 n \n".encode()
    trailer = f"trailer\n<< /Size {n} /Root 1 0 R /Info 5 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()

    rev1 = bytes(body) + bytes(xref) + trailer
    rev1_startxref = xref_offset

    # --- incremental update: revised Info object (object 5) only ---
    upd = bytearray()
    upd_base = len(rev1)
    upd_obj5_offset = upd_base + len(upd)
    upd += f"5 0 obj\n<< /Author ({rev2_author}) >>\nendobj\n".encode()

    upd_xref_offset = upd_base + len(upd)
    upd += b"xref\n0 1\n0000000000 65535 f \n5 1\n"
    upd += f"{upd_obj5_offset:010d} 00000 n \n".encode()
    upd += (
        f"trailer\n<< /Size 6 /Root 1 0 R /Info 5 0 R /Prev {rev1_startxref} >>\n"
        f"startxref\n{upd_xref_offset}\n%%EOF\n"
    ).encode()

    out = rev1 + bytes(upd)
    path.write_bytes(out)
    return out


def make_pdf_signed(path: Path, *, author: str = "Signer Author") -> bytes:
    """A minimal /AcroForm with SigFlags bit 1 set — enough to exercise
    fmt_pdf._is_signed()'s detection without a real cryptographic signature
    (irrelevant to the tier-refusal logic under test, which reads only the
    flag).
    """
    import pikepdf

    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.docinfo["/Author"] = author
    pdf.Root.AcroForm = pdf.make_indirect(pikepdf.Dictionary(SigFlags=3, Fields=pikepdf.Array([])))
    pdf.save(path)
    return path.read_bytes()


def make_pdf_encrypted(path: Path, *, user_password: str = "secret") -> bytes:
    import pikepdf

    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.save(path, encryption=pikepdf.Encryption(user=user_password, owner="owner-pw"))
    return path.read_bytes()


def _attach_file(pdf, name: str, data: bytes, *, relationship: str | None):
    import pikepdf

    stream = pdf.make_stream(data)
    stream["/Type"] = pikepdf.Name("/EmbeddedFile")
    ef = pikepdf.Dictionary(F=stream)
    fs = pikepdf.Dictionary(Type=pikepdf.Name("/Filespec"), F=name, UF=name, EF=ef)
    if relationship is not None:
        fs["/AFRelationship"] = pikepdf.Name(relationship)
    return pdf.make_indirect(fs)


def make_pdf_with_embedded_attachment(
    path: Path, *, name: str = "factur-x.xml", data: bytes = b"<Invoice>real content</Invoice>"
) -> bytes:
    """A PDF carrying a legitimate, non-C2PA embedded file attachment (the
    Factur-X/ZUGFeRD case): reachable both from Root/AF and from the
    /Names/EmbeddedFiles name tree, exactly as a real hybrid invoice is.
    """
    import pikepdf

    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    pdf.docinfo["/Author"] = "Attachment Author"
    fs = _attach_file(pdf, name, data, relationship="/Data")
    pdf.Root.AF = pikepdf.Array([fs])
    pdf.Root.Names = pdf.make_indirect(
        pikepdf.Dictionary(EmbeddedFiles=pikepdf.Dictionary(Names=pikepdf.Array([name, fs])))
    )
    pdf.save(path)
    return path.read_bytes()


def make_pdf_with_c2pa_and_attachment(path: Path) -> bytes:
    """Both a C2PA manifest attachment and a legitimate one, in both /AF and
    the /Names/EmbeddedFiles tree — the fixture that separates "prune the
    manifest" from "destroy every attachment".
    """
    import pikepdf

    pdf = pikepdf.new()
    pdf.add_blank_page(page_size=(200, 200))
    manifest = _attach_file(pdf, "c2pa.manifest", b"fake-c2pa-manifest", relationship="/C2PA_Manifest")
    invoice = _attach_file(pdf, "factur-x.xml", b"<Invoice>real content</Invoice>", relationship="/Data")
    pdf.Root.AF = pikepdf.Array([manifest, invoice])
    pdf.Root.Names = pdf.make_indirect(
        pikepdf.Dictionary(
            EmbeddedFiles=pikepdf.Dictionary(
                Names=pikepdf.Array(["c2pa.manifest", manifest, "factur-x.xml", invoice])
            )
        )
    )
    pdf.save(path)
    return path.read_bytes()


def make_jpeg_with_xmp(path: Path, *, creator: str = "XMP Creator") -> bytes:
    """A JPEG whose ONLY identity metadata lives in an XMP APP1 packet — no
    EXIF, no IPTC. exiftool reports it under namespace-qualified `-G1` groups
    (XMP-dc, XMP-x), never the bare group name "XMP".
    """
    from PIL import Image

    img = Image.new("RGB", (16, 16), color=(9, 9, 9))
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    base = buf.getvalue()

    packet = (
        f'<?xpacket begin="{chr(0xFEFF)}" id="W5M0MpCehiHzreSzNTczkc9d"?>'
        '<x:xmpmeta xmlns:x="adobe:ns:meta/">'
        '<rdf:RDF xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">'
        '<rdf:Description rdf:about="" xmlns:dc="http://purl.org/dc/elements/1.1/">'
        f"<dc:creator><rdf:Seq><rdf:li>{creator}</rdf:li></rdf:Seq></dc:creator>"
        "</rdf:Description></rdf:RDF></x:xmpmeta>"
        '<?xpacket end="w"?>'
    ).encode("utf-8")
    payload = b"http://ns.adobe.com/xap/1.0/\x00" + packet
    seg = b"\xff\xe1" + struct.pack(">H", len(payload) + 2) + payload

    out = base[:2] + seg + base[2:]
    path.write_bytes(out)
    return out


def make_tiff_with_exif(path: Path, *, author: str = "Tiff Author") -> bytes:
    from PIL import Image

    img = Image.new("RGB", (16, 16), color=(7, 7, 7))
    img.save(path, format="TIFF", description=author)
    return path.read_bytes()


def make_utf16_text_with_zwsp(path: Path, *, big_endian: bool = False) -> bytes:
    """A BOM-prefixed UTF-16 text file containing a zero-width space —
    watermark_scan._decode supports it, so watermark_remove must write it back
    in the SAME encoding rather than silently transcoding it to UTF-8.
    """
    text = "hello" + chr(0x200B) + "world\n"
    if big_endian:
        raw = b"\xfe\xff" + text.encode("utf-16-be")
    else:
        raw = b"\xff\xfe" + text.encode("utf-16-le")
    path.write_bytes(raw)
    return raw


# ----------------------------------------------------------------- ooxml ---


def make_docx_simple(path: Path, *, author: str = "Docx Author", company: str = "Acme") -> bytes:
    import docx

    doc = docx.Document()
    doc.add_paragraph("hello world")
    doc.core_properties.author = author
    doc.core_properties.last_modified_by = author
    doc.save(path)

    # python-docx doesn't set Company by default — patch it in via zip surgery
    entries = _read_zip(path)
    if "docProps/app.xml" in entries:
        from lxml import etree

        root = etree.fromstring(entries["docProps/app.xml"])
        ns = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
        el = etree.SubElement(root, f"{{{ns}}}Company")
        el.text = company
        entries["docProps/app.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _write_zip(path, entries)
    return path.read_bytes()


def make_docx_with_thumbnail_only_jpeg(path: Path) -> bytes:
    """docProps/thumbnail.jpeg is the ONLY part with a .jpeg extension —
    after removal, the Content_Types Default Extension="jpeg" entry (if
    the builder emitted one) must survive since Defaults with zero
    matching parts are legal OPC, never removed just because their last
    matching part disappeared.
    """
    import docx
    from lxml import etree

    doc = docx.Document()
    doc.add_paragraph("hello")
    doc.save(path)

    entries = _read_zip(path)
    # minimal 2x2 JPEG bytes as a stand-in thumbnail
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (2, 2)).save(buf, format="JPEG")
    entries["docProps/thumbnail.jpeg"] = buf.getvalue()

    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ct_root = etree.fromstring(entries["[Content_Types].xml"])
    if not any(e.get("Extension", "").lower() == "jpeg" for e in ct_root.findall(f"{{{ct_ns}}}Default")):
        default = etree.SubElement(ct_root, f"{{{ct_ns}}}Default")
        default.set("Extension", "jpeg")
        default.set("ContentType", "image/jpeg")
    entries["[Content_Types].xml"] = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    rels_root = etree.fromstring(entries["_rels/.rels"])
    rel = etree.SubElement(rels_root, f"{{{rel_ns}}}Relationship")
    rel.set("Id", "rIdThumb")
    rel.set("Type", "http://schemas.openxmlformats.org/package/2006/relationships/metadata/thumbnail")
    rel.set("Target", "docProps/thumbnail.jpeg")
    entries["_rels/.rels"] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    _write_zip(path, entries)
    return path.read_bytes()


def make_docx_with_dangling_relationship(path: Path) -> bytes:
    """A pre-existing dangling relationship the tool never touches: a
    package-root relationship pointing at a part that does not exist in
    the package, present in the INPUT already — not something the tool's
    own removal introduces. The tool must not "fix" it or treat its own
    output as newly broken.
    """
    import docx
    from lxml import etree

    doc = docx.Document()
    doc.add_paragraph("hello")
    doc.save(path)

    entries = _read_zip(path)
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    rels_root = etree.fromstring(entries["_rels/.rels"])
    rel = etree.SubElement(rels_root, f"{{{rel_ns}}}Relationship")
    rel.set("Id", "rIdDangling")
    rel.set("Type", "http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties-extra")
    rel.set("Target", "docProps/does-not-exist.xml")
    entries["_rels/.rels"] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _write_zip(path, entries)
    return path.read_bytes()


def make_docx_with_external_attached_template(path: Path) -> bytes:
    import docx
    from lxml import etree

    doc = docx.Document()
    doc.add_paragraph("hello")
    doc.save(path)

    entries = _read_zip(path)
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    settings_root = etree.fromstring(entries["word/settings.xml"])
    attached = etree.SubElement(settings_root, f"{{{w_ns}}}attachedTemplate")
    attached.set(f"{{{r_ns}}}id", "rIdTemplate")
    entries["word/settings.xml"] = etree.tostring(
        settings_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    settings_rels_path = "word/_rels/settings.xml.rels"
    if settings_rels_path in entries:
        rels_root = etree.fromstring(entries[settings_rels_path])
    else:
        rels_root = etree.Element(f"{{{rel_ns}}}Relationships", nsmap={None: rel_ns})
    rel = etree.SubElement(rels_root, f"{{{rel_ns}}}Relationship")
    rel.set("Id", "rIdTemplate")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate")
    rel.set("Target", r"\\server\share\Normal.dotm")
    rel.set("TargetMode", "External")
    entries[settings_rels_path] = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    _write_zip(path, entries)
    return path.read_bytes()


def make_zip_bomb(path: Path) -> bytes:
    """A single member with a wildly inflated declared uncompressed size vs.
    actual compressed data — enough to trip a decompression-ratio guard
    without actually allocating gigabytes to build the fixture.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _DOCX_CONTENT_TYPES)
        info = zipfile.ZipInfo("bomb.bin")
        info.compress_type = zipfile.ZIP_DEFLATED
        zf.writestr(info, b"\x00" * (50 * 1024 * 1024))
    raw = buf.getvalue()
    path.write_bytes(raw)
    return raw


def make_zip_with_path_traversal(path: Path) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _DOCX_CONTENT_TYPES)
        zf.writestr("../../etc/evil.xml", b"<evil/>")
    raw = buf.getvalue()
    path.write_bytes(raw)
    return raw


def make_zip_with_duplicate_entry_names(path: Path) -> bytes:
    """Two entries sharing one name — zipfile.namelist()/read() silently
    resolve to the last one, masking a polyglot-style duplicate-entry attack
    unless explicitly rejected before any member is read.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _DOCX_CONTENT_TYPES)
        zf.writestr("word/document.xml", b"<first/>")
        zf.writestr("word/document.xml", b"<second/>")
    raw = buf.getvalue()
    path.write_bytes(raw)
    return raw


# --------------------------------------------------------------- helpers ---


def _read_zip(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as zf:
        return {n: zf.read(n) for n in zf.namelist()}


def _write_zip(path: Path, entries: dict[str, bytes]) -> None:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in entries.items():
            zf.writestr(name, data)
    path.write_bytes(buf.getvalue())
